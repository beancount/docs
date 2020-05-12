import argparse
import hashlib
import io
import os
import shutil

import requests
import pypandoc
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.packuri import PackURI
from docx.parts.image import ImagePart

INTERMEDIATE_FMT = 'docx'


def download_document(
    document_id: str,
    fmt=INTERMEDIATE_FMT,
    file_name=None,
) -> bytes:
    """
    Export and download. Formats available: html, docx, odt, rtf, txt
    https://googlesystem.blogspot.com/2007/07/download-published-documents-and.html
    """
    url = f'https://docs.google.com/document/export?format={fmt}&id={document_id}'
    response = requests.get(url)
    response.raise_for_status()
    data = response.content
    if file_name:
        with open(file_name, 'wb') as file:
            file.write(data)
    return data


def download_drawings(document_id: str, drawing_dir: str):
    document_html = download_document(document_id, fmt='html')
    soup = BeautifulSoup(document_html, 'html.parser')
    os.makedirs(drawing_dir, exist_ok=True)
    for idx, image in enumerate(soup.select('img[src*="drawings"]')):
        response = requests.get(image['src'])
        response.raise_for_status()
        file_name = os.path.join(drawing_dir, f'{idx:02d}.png')
        with open(file_name, 'wb') as file:
            file.write(response.content)


def prepare_docx(file_name: str, drawing_dir: str = None) -> bytes:
    """
    Prepare docx document for Pandoc conversion:
    * Mark code blocks with SourceCode style
    * Replace vector graphics with raster
    """
    doc = Document(file_name)
    doc.styles.add_style('SourceCode', WD_STYLE_TYPE.PARAGRAPH)
    drawing_idx = 0

    for para in doc.paragraphs:
        if len(para.runs) == 0:
            continue

        for run_idx, run in enumerate(para.runs):
            if run_idx == 0 and run.text == '\t':
                continue  # Ignore leading tabs
            if run_idx == 0 and run.font.name in ['Consolas', 'Courier New']:
                # If paragraph starts with a snippet in monospace font,
                # consider it a code block and mark it with SourceCode style
                # https://groups.google.com/d/msg/pandoc-discuss/SIwE9dhGF4U/Wjy8zmQ1CQAJ
                para.style = doc.styles['SourceCode']
                continue
            if run.font.name in ['Consolas', 'Courier New']:
                # Mark with striketrough style to convert to inline code later
                run.font.strike = True

        if para.runs[0].element.xpath(
            './/*[@uri="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"]'
        ):
            # WordprocessingML group found
            if drawing_dir:
                # Pandoc can't convert embedded vector graphics
                # So we insert previously downloaded image in the same run
                drawing_path = os.path.join(drawing_dir,
                                            f'{drawing_idx:02d}.png')
                para.runs[0].add_picture(drawing_path)
                drawing_idx += 1

    for part in doc.part.related_parts.values():
        if isinstance(part, ImagePart):
            # Use deterministic names for images
            image_hash = part.image.sha1
            image_name = os.path.join(
                part.partname.baseURI,
                f'{image_hash}.{part.partname.ext}',
            )
            part.partname = PackURI(image_name)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def convert(
    document: bytes,
    output: str,
    input_fmt: str = INTERMEDIATE_FMT,
    output_fmt: str = 'markdown_strict',
):
    # Enable footnotes extension
    if output_fmt.startswith('markdown'):
        output_fmt += '+footnotes'

    output_dir = os.path.dirname(output)
    image_dir = os.path.splitext(output)[0]
    # Directory path for --extract-media must be relative to the output dir
    image_dir_rel = os.path.relpath(image_dir, output_dir)
    pypandoc.convert_text(
        document,
        output_fmt,
        format=input_fmt,
        extra_args=[
            '--wrap=none',  # Don't wrap lines
            '--standalone',  # Don't strip headers
            '--extract-media', image_dir_rel,
        ],
        filters=[
            'filter.py',
        ],
        outputfile=output,
    )
    if output_dir and os.path.exists(image_dir_rel):
        # Move image dir to output dir
        shutil.rmtree(image_dir, ignore_errors=True)
        shutil.move(image_dir_rel, image_dir)


def main():
    parser = argparse.ArgumentParser()
    subparsers = parser.add_subparsers(dest='command')
    document_parser = subparsers.add_parser(
        'document',
        help='download document')
    document_parser.add_argument('id', help='document id')
    document_parser.add_argument('output', help='output file name')

    drawings_parser = subparsers.add_parser(
        'drawings',
        help='download drawings')
    drawings_parser.add_argument('id', help='document id')
    drawings_parser.add_argument('dir', help='output directory name')

    convert_parser = subparsers.add_parser(
        'convert',
        help='convert document')
    convert_parser.add_argument('name', help='docx file name')
    convert_parser.add_argument('output', help='output file name')
    convert_parser.add_argument(
        '--format',
        help='output format',
        default='markdown_strict')
    convert_parser.add_argument(
        '--drawing-dir',
        help='drawing directory')
    args = parser.parse_args()

    if args.command == 'document':
        download_document(args.id, file_name=args.output)
    elif args.command == 'drawings':
        download_drawings(args.id, args.dir)
    elif args.command == 'convert':
        document_path = args.name
        document = prepare_docx(
            document_path,
            drawing_dir=args.drawing_dir,
        )
        convert(
            document,
            args.output,
            output_fmt=args.format,
        )


if __name__ == '__main__':
    main()
